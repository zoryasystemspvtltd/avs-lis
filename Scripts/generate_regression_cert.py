# -*- coding: utf-8 -*-
"""Regression certification report for AVILIS."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

OUTPUT = r"I:\Projects\LIS\regression-certification.docx"


def build():
    doc = Document()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("AVILIS / ZoryaLIS\nFull Regression & Stability Certification")
    r.bold = True
    r.font.size = Pt(18)
    doc.add_paragraph(f"Date: {datetime.now():%Y-%m-%d %H:%M}")
    doc.add_paragraph("Environment: PRE-PRODUCTION — I:\\Projects\\PROD\\AVILIS\\")
    doc.add_page_break()

    sections = [
        ("1. Executive Assessment", [
            "CERTIFICATION STATUS: CONDITIONAL — NOT READY FOR PRODUCTION",
            "Portal and API binaries built and deployed. API IIS site was found STOPPED post-deploy (remediated).",
            "SQL Server database AVSLIS returned OS error 21 (device not ready) on MDF read — blocks full data regression.",
            "Authenticated API calls return HTTP 200 with null body for master list endpoints when DB fails (exception swallowed).",
        ]),
        ("2. Regression Coverage Executed", [
            "API Release build (MSBuild) — PASS",
            "Angular production build — PASS (budget warnings only)",
            "Deploy API + Portal to PROD paths — PASS",
            "IIS restart — PASS",
            "Portal HTTP 200 — PASS",
            "API root HTTP 200 — PASS (slow cold start ~4–16s)",
            "OAuth /Token login — PASS (slow)",
            "api/Lis ping — PASS (returns false without equipment key)",
            "SQL integrity queries — BLOCKED (DB IO error)",
            "Full UI click-through all masters — NOT COMPLETED (DB dependency)",
        ]),
        ("3. Critical Findings (P0)", [
            "P0-1: AVSLIS database unavailable — Sqlcmd error 823, MDF path I:\\Projects\\SQL-DATA\\AVSLIS.mdf device not ready.",
            "P0-2: AVILIS_API IIS website was Stopped after deployment — lists/masters fail until site started.",
            "P0-3: API master GET endpoints return literal null on DB exception (e.g. DepartmentController) — UI shows empty grids instead of explicit error.",
        ]),
        ("4. Issues Fixed This Pass", [
            "Test Master Edit: PascalCase/camelCase binding for specimen, department, isActive; dropdown Code/Name fallbacks; loading flag on success.",
            "build-portal.bat: fixed PowerShell here-string failure; new patch-portal-env.ps1 ensures ApplicationServer=http://localhost:8081 in prod bundle.",
            "deploy-api.bat: added ensure-api-site-started.ps1 to start AVILIS_API site and app pool after deploy.",
        ]),
        ("5. Previously Fixed (Revalidated by Code Review)", [
            "Department search (case-insensitive), create validation, delete button",
            "Unit/Method create IsActive defaults",
            "Specimen list sort column fix, create validation",
            "Test create auto-code, getRawValue for disabled fields",
            "Test edit isActive coerceBool",
            "Test mapping auto-populate, parameter range validation",
            "Sale invoice patient required (UI + SaleInvoiceManager)",
            "Sale invoice billable tests filtered by active rate effective dates",
            "Recent sample TestParameterNames column",
            "Referral Doctor menu hidden",
            "Diagnostic report module + workflow gates",
            "QAuthorize 401 for unauthenticated users",
        ]),
        ("6. Open / Deferred Issues", [
            "P1: API response time 15+ seconds on cold start — investigate DB connectivity and EF initialization",
            "P1: TestReportLabNumbers returns 401 Insufficient privilege for standard admin token — verify Reports module permission mapping",
            "P2: Many controllers catch exceptions and return null (masks failures)",
            "P2: LisController/ResultsController AllowAnonymous — equipment AccessKey only",
            "P2: EquipmentParamMappingManager not implemented",
            "P2: Angular initial bundle exceeds 2MB budget",
            "P3: Portal dist folder retains older main-*.js files (cosmetic)",
        ]),
        ("7. Module Regression Matrix (Status)", [
            "Authentication: Code review PASS; runtime Token PASS (slow); full logout/session expiry not automated",
            "Master CRUD: BLOCKED for runtime — DB; code paths reviewed PASS for prior QA fixes",
            "Sale Invoice: Code review PASS (patient, rates, link test requests)",
            "Reports: Code review PASS; TestReport API auth issue noted",
            "Analyzer flow: Documentation PASS (prior analysis); runtime not executed",
            "Approval workflow: Code review PASS (ReportStatusType 0–7)",
        ]),
        ("8. Build & Deployment", [
            "API: Lis.Api.dll Release build succeeded",
            "Portal: main-es2015.757fb328142945680304.js — contains localhost:8081",
            "Deployed: I:\\Projects\\PROD\\AVILIS\\API and PORTAL",
            "IIS: Restarted; AVILIS_API started via remediation script",
        ]),
        ("9. UAT Readiness", [
            "DO NOT promote to production until AVSLIS database is online and DBCC CHECKDB passes.",
            "Re-run full master CRUD, invoice, approval, and report UAT after DB restore.",
            "Verify Reports module permission for diagnostic report role.",
            "Confirm AVILIS_API site auto-starts after server reboot (IIS site startup setting).",
        ]),
        ("10. Recommended Next Steps", [
            "1. Restore/re-attach AVSLIS.mdf on I:\\Projects\\SQL-DATA or update connection string",
            "2. Run DBCC CHECKDB on AVSLIS",
            "3. Execute full manual UAT script against http://localhost:8080",
            "4. Fix DepartmentController (and peers) to return 500 or structured error instead of null",
            "5. Assign Reports CanView to roles needing Diagnostic Report",
        ]),
    ]

    for title, items in sections:
        doc.add_heading(title, level=1)
        for item in items:
            doc.add_paragraph(item, style="List Bullet" if not item.startswith("CERTIFICATION") else None)
        doc.add_paragraph()

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
