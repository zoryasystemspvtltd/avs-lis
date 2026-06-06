# -*- coding: utf-8 -*-
"""Generate AVILIS discovery document from implemented codebase facts."""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"I:\Projects\LIS\feature.docx"


def set_cell_shading(cell, fill="E8F4F6"):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "1A7F8C")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        if level:
            p.paragraph_format.left_indent = Inches(0.25 * level)


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def build_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Cover
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("AVILIS")
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = RGBColor(26, 127, 140)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Laboratory Information System")
    r2.font.size = Pt(18)
    r2.font.color.rgb = RGBColor(80, 80, 80)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = sub2.add_run("Product, Technical, Architecture, Workflow & Capability Discovery Report")
    r3.font.size = Pt(14)
    r3.bold = True

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Source: AVS-LIS Implemented Codebase\n").font.size = Pt(11)
    meta.add_run("Repository: I:\\Projects\\LIS\\avs-lis\n").font.size = Pt(11)
    meta.add_run("Document Type: Master Source for Product & Technical Presentations\n").font.size = Pt(11)
    meta.add_run("Classification: Implementation Discovery (Facts from Code Only)").font.size = Pt(11)

    doc.add_page_break()

    # TOC placeholder
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Product Overview",
        "2. Technical Architecture",
        "3. Module Inventory",
        "4. Business Workflows",
        "5. Test & Billing Engine",
        "6. Analyzer & Integration Capability",
        "7. Reporting Capability",
        "8. Security & Access Control",
        "9. UI/UX & Technical Strengths",
        "10. Deployment & Scalability",
        "11. Commercial Positioning Inputs",
        "12. Conclusion",
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    doc.add_page_break()

    # 1 Product Overview
    doc.add_heading("1. Product Overview", level=1)
    doc.add_paragraph(
        "AVILIS (AVS-LIS) is a web-based Laboratory Information System implemented as a dual-tier "
        "application: an Angular 10 portal and a .NET Framework 4.8 Web API, backed by Microsoft SQL Server "
        "(database AVSLIS). The product supports diagnostic laboratory and hospital laboratory operations "
        "including master data management, patient registration, sale invoicing, sample workflow, analyzer "
        "result ingestion, technician and doctor approval, operational reporting, and final diagnostic report printing."
    )
    doc.add_heading("1.1 Purpose and Operational Scope", level=2)
    add_bullets(doc, [
        "Centralize pathology and clinical laboratory test catalog, pricing, and specimen configuration.",
        "Manage patient demographics and laboratory test bookings linked to sale invoices.",
        "Track samples from collection through equipment transmission, result entry, and authorization.",
        "Support billing via sale invoices with corporate, referral doctor, profile, and emergency rate types.",
        "Provide operational registers and printable final diagnostic test reports for paid, approved cases.",
        "Integrate with clinical analyzers via HL7 over TCP/IP (Metis6000 implementation) and optional HIS order intake.",
    ])
    doc.add_heading("1.2 Target Usage Context", level=2)
    add_bullets(doc, [
        "Standalone diagnostic centers and hospital laboratory departments.",
        "Environments requiring instrument bidirectional communication (orders/results) with LIS middleware.",
        "Facilities needing configurable test rates, corporate contracts, and referral doctor pricing.",
        "Organizations deploying on-premises under IIS with separated API and portal hosts.",
    ])
    doc.add_heading("1.3 Patient and Test Lifecycle (Implemented)", level=2)
    add_numbered(doc, [
        "Patient registered in Patient Master (PatientDetails table).",
        "Sale Invoice created with test lines; rates resolved by invoice date and priority rules.",
        "Test booking / sample request created (TestRequestDetails) with HISRequestNo as lab/request identifier.",
        "Sample collected; status progresses through ReportStatusType workflow.",
        "Analyzer or manual result entry populates TestResults and TestResultDetails.",
        "Technician review and Doctor authorization update ReportStatus and authorization metadata.",
        "Payment completed on invoice (PaymentStatusType.Paid).",
        "Final Diagnostic Report printed from Reports module when all gates pass.",
    ])

    # 2 Architecture
    doc.add_heading("2. Technical Architecture", level=1)
    doc.add_heading("2.1 Solution Structure", level=2)
    add_table(doc, ["Layer", "Project / Path", "Technology"], [
        ["Presentation", "web/Lis.Web", "Angular 10.0.5, Bootstrap 3.4, ng-select, Zorya theme CSS"],
        ["API", "web/Lis.Api", "ASP.NET Web API 5.2, OWIN OAuth Bearer"],
        ["Business Logic", "LIS.Businesslogic", "Manager pattern, ModuleRepo, GenericUnitOfWork"],
        ["Data Access", "LIS.DataModel (LIS.DataAccess)", "Entity Framework 6.4"],
        ["DTO / Entities", "LIS.DtoModel", "EF-mapped entities and report DTOs"],
        ["Analyzer Middleware", "LIS.Com.Businesslogic, LisTCPIPConsole", "HL7 TCP listener, REST client to API"],
        ["Barcode Utility", "BarcodePrint", "WinForms label printing via API"],
        ["Database", "AVSLIS on SQL Server", "SQLEXPRESS (deployed config)"],
    ], [1.2, 2.2, 1.6])

    doc.add_heading("2.2 Frontend Architecture", level=2)
    add_bullets(doc, [
        "Single-page application with RouterModule; AuthGuard on protected routes.",
        "Schema-driven master CRUD: MasterListComponent + MasterFormComponent + MASTER_LIST_SCHEMAS.",
        "Generic list grid: app-list-module with ModuleService/MasterService and ApiOption headers.",
        "Operational reports extend ReportPageBase (date range max 2 months, Excel export).",
        "Diagnostic Report uses dedicated TestReportComponent with Lab No dropdown.",
        "OAuth token stored in localStorage; permissions drive left-nav visibility.",
        "Print: window.print() with body CSS classes; ngx-print on sample lists; sale invoice print route.",
    ])
    doc.add_heading("2.3 Backend Architecture", level=2)
    add_bullets(doc, [
        "Attribute routing (api/Reports, api/SaleInvoice, api/TestRate, etc.) plus convention routes.",
        "SimpleInjector DI with scoped managers and GenericUnitOfWork per request.",
        "QAuthorize attribute enforces module-level bitwise permissions on controllers.",
        "MasterApiControllerBase pattern for standardized master CRUD endpoints.",
        "SignalR hub mapped at /lis for equipment heartbeat notifications.",
    ])
    doc.add_heading("2.4 Database Architecture", level=2)
    add_bullets(doc, [
        "Entity Framework 6 with [Table] attributes on DTO model classes.",
        "Master tables: ReferralDoctor, Corporate, TestGroup, TestCategory, Unit, Method, etc.",
        "HIS catalog: HISTestMaster, HISParameterMaster, HISParameterRangMaster, HISSpecimenMaster.",
        "Transaction: PatientDetails, TestRequestDetails, SaleInvoice, SaleInvoiceDetail.",
        "Results: TestResults, TestResultDetails (+ history tables TestResultsHist).",
        "Connection: Server=.\\SQLEXPRESS; Database=AVSLIS; Trusted_Connection (Web.config).",
    ])
    doc.add_heading("2.5 Deployment Architecture", level=2)
    add_table(doc, ["Component", "Deploy Path", "Port / URL"], [
        ["API (IIS)", "I:\\Projects\\PROD\\AVILIS\\API", "http://localhost:8081"],
        ["Portal (IIS static)", "I:\\Projects\\PROD\\AVILIS\\PORTAL", "http://localhost:8080"],
        ["Build scripts", "I:\\Projects\\PROD\\AVILIS\\SCRIPTS", "deploy-api.bat, deploy-portal.bat"],
        ["IIS Sites", "AVILIS_API, AVILIS_PORTAL", "Separate app pools"],
    ], [1.4, 2.4, 1.2])

    # 3 Modules
    doc.add_heading("3. Module Inventory", level=1)
    doc.add_paragraph("All modules below are implemented with Angular routes, API endpoints, and business managers unless noted.")

    modules_setup = [
        ("Department", "Department master", "DepartmentManager", "Setup menu", "Implemented"),
        ("Unit", "Result/report units", "UnitManager", "Setup", "Implemented"),
        ("Method", "Analytical methods", "MethodManager", "Setup", "Implemented"),
        ("Sample Type", "Sample type classification", "SampleTypeManager", "Setup", "Implemented"),
        ("Container", "Collection container types", "ContainerManager", "Setup", "Implemented"),
        ("Equipment", "Analyzer registration, AccessKey", "EquipmentManager", "Setup", "Implemented"),
        ("Equipment Heartbeat", "Live equipment status", "EquipmentHeartBeatManager", "Setup", "Implemented"),
    ]
    doc.add_heading("3.1 Setup Modules", level=2)
    add_table(doc, ["Module", "Purpose", "Backend", "UI Access", "Status"], modules_setup, [1.1, 1.8, 1.0, 0.9, 0.7])

    modules_master = [
        ("Test Master (HisTest)", "HIS test catalog", "HISTestMasterManager", "Master", "Implemented"),
        ("Specimen", "Specimen types", "SpecimenManager", "Master", "Implemented"),
        ("Test Group / Category", "Test grouping", "TestGroupManager, TestCategoryManager", "Master", "Implemented"),
        ("Test Profile", "Package tests with details", "TestProfileMasterManager", "Master", "Implemented"),
        ("Test Rate", "Effective-dated pricing", "TestRateMasterManager", "Master", "Implemented"),
        ("Referral Doctor", "Referring physicians", "ReferralDoctorManager", "Master", "Implemented"),
        ("Corporate", "Corporate clients/discounts", "CorporateManager", "Master", "Implemented"),
        ("Test Mapping", "Equipment-HIS-LIS test map", "TestMappingCrudManager", "Master", "Implemented"),
        ("Test Parameter", "Catalog parameters", "TestParameterCatalogManager", "Master", "Implemented"),
        ("Parameter Master", "HIS parameters per test", "HisParameterMasterManager", "Master", "Implemented"),
        ("Parameter Range", "Reference ranges", "HisParameterRangeCrudManager", "Master", "Implemented"),
        ("Patient Master", "Patient demographics", "PatientMasterManager", "Transaction", "Implemented"),
        ("Sale Invoice", "Billing and booking", "SaleInvoiceManager", "Transaction", "Implemented"),
    ]
    doc.add_heading("3.2 Master & Transaction Modules", level=2)
    add_table(doc, ["Module", "Purpose", "Backend", "Menu", "Status"], modules_master, [1.0, 1.6, 1.2, 0.8, 0.7])

    modules_workflow = [
        ("Recent Samples", "New/pending test requests", "TestRequestDetailsManager", "Working Board", "Implemented"),
        ("Technician Approval", "Tech review of results", "SampleController POST", "Working Board", "Implemented"),
        ("Doctor Approval", "Doctor authorization", "SampleController PUT", "Working Board", "Implemented"),
        ("Approved / Rejected Samples", "Post-review lists", "TestRequestDetailsManager", "Working Board", "Implemented"),
        ("Quality Controls", "QC sample management", "QualityControlManager", "Working Board", "Implemented"),
        ("Users / Roles", "Identity and RBAC", "UsersController, RolesController", "Account", "Implemented"),
    ]
    doc.add_heading("3.3 Workflow & Security Modules", level=2)
    add_table(doc, ["Module", "Purpose", "Backend", "Menu", "Status"], modules_workflow, [1.1, 1.7, 1.1, 0.9, 0.7])

    doc.add_heading("3.4 Reports", level=2)
    add_table(doc, ["Report", "Purpose", "API Endpoint", "Export/Print", "Status"], [
        ("Sale Invoice Register", "Filtered invoice register", "GET api/Reports/SaleInvoiceRegister", "Excel export", "Implemented"),
        ("Test Booking Register", "Test booking register", "GET api/Reports/TestBookingRegister", "Excel export", "Implemented"),
        ("Diagnostic Report", "Final diagnostic test report", "GET api/Reports/TestReport, TestReportLabNumbers", "Browser print A4", "Implemented"),
        ("Sale Invoice Print", "Invoice print layout", "UI route sale-invoices/print/:id", "window.print", "Implemented"),
    ], [1.3, 1.5, 1.8, 1.0, 0.7])

    # 4 Workflows
    doc.add_heading("4. Business Workflows", level=1)
    workflows = [
        ("Patient Registration", "Patient Master CRUD", "PatientDetails → api/PatientMaster", "Required for sale invoice"),
        ("Sale Invoice", "Create invoice, add tests, apply rates", "SaleInvoiceManager, TestRate GetEffectiveForInvoice", "Creates/links TestRequestDetails; patient mandatory"),
        ("Sample Collection", "HIS order or invoice-driven booking", "PatientDetailManager.CreateNewOrder, NewSampleController", "HISRequestNo assigned"),
        ("Result Entry", "Analyzer POST or manual", "ResultManager.SaveTestResult → TestResults/Details", "Status → ReportGenerated"),
        ("Technician Approval", "Review results", "TestRequestDetailsManager.TechnicianReview", "ReportStatus TechnicianApproved/Rejected"),
        ("Doctor Authorization", "Final clinical approval", "TestRequestDetailsManager.DoctorReview", "ReportStatus DoctorApproved; AuthorizationDate set"),
        ("Payment", "Invoice payment status", "SaleInvoiceManager status API", "PaymentStatusType.Paid required for diagnostic report"),
        ("Diagnostic Report Print", "Search by Lab No dropdown", "TestReportManager validations", "Paid + DoctorApproved + results"),
    ]
    add_table(doc, ["Workflow", "Description", "Key Components", "Gate/Output"], workflows, [1.1, 1.5, 1.8, 1.2])

    doc.add_heading("4.1 Report Status Workflow (ReportStatusType)", level=2)
    add_table(doc, ["Value", "Status", "Meaning"], [
        ("0", "New", "Initial request"),
        ("1", "SentToEquipment", "Transmitted to analyzer"),
        ("2", "ReportGenerated", "Results recorded"),
        ("3", "TechnicianApproved", "Technician signed off"),
        ("4", "TechnicianRejected", "Rejected at tech level"),
        ("5", "DoctorApproved", "Authorized for release"),
        ("6", "DoctorRejected", "Rejected at doctor level"),
        ("7", "FinallyRejected", "Terminal rejection"),
    ], [0.6, 1.2, 2.6])

    # 5 Billing
    doc.add_heading("5. Test & Billing Engine", level=1)
    doc.add_paragraph(
        "TestRateMasterManager implements effective-dated rate resolution. SaleInvoiceManager recalculates "
        "line amounts using GetEffectiveRateForInvoice when rate is not manually set."
    )
    doc.add_heading("5.1 Rate Type Priority (Invoice Date)", level=2)
    add_numbered(doc, [
        "Emergency (if useEmergency flag) — uses EmergencyRate when > 0",
        "Corporate (CorporateId on invoice)",
        "Referral Doctor (ReferralDoctorId)",
        "Test Profile (profileId when applicable)",
        "Standard (fallback)",
    ])
    doc.add_heading("5.2 Invoice Model", level=2)
    add_bullets(doc, [
        "InvoiceStatusType: Draft, Confirmed, Paid, Cancelled",
        "PaymentStatusType: Unpaid, Partial, Paid",
        "SaleInvoiceDetail links TestId, Rate, Quantity, Discount, Tax, NetAmount, RequestDetailId, SampleNo",
        "Next invoice number via api/SaleInvoice/NextInvoiceNo",
        "Cancel via PUT api/SaleInvoice/Cancel/{id}",
    ])

    # 6 Analyzer
    doc.add_heading("6. Analyzer & Integration Capability", level=1)
    add_table(doc, ["Capability", "Implementation", "Status"], [
        ("HL7 over TCP/IP", "TCPIPHL7Command, Metis6000TCPIPCommand", "Implemented (Metis6000)"),
        ("ASTM over TCP/IP", "TCPIPASTMCommand", "Listener exists; InitTCPIPCommand throws NotImplemented"),
        ("Equipment type enum", "EquipmentType.Metis6000 only", "Implemented"),
        ("Result import to LIS", "POST api/Lis → ResultManager", "Implemented (AllowAnonymous)"),
        ("Order/query to analyzer", "GET api/Lis/{sampleNo}, panel check", "Implemented"),
        ("Equipment heartbeat", "POST api/Heartbeat, SignalR /lis hub", "Implemented"),
        ("Equipment test mapping", "EquipmentTestMappingsController", "Implemented"),
        ("Equipment param mapping", "EquipmentParamMappingsController", "Implemented"),
        ("Barcode printing", "BarcodePrint WinForms + api/NewSample", "Implemented (desktop)"),
        ("LisTCPIPConsole", "Tray app bridging analyzer to API", "Implemented (separate deploy)"),
    ], [1.4, 2.8, 0.9])

    doc.add_heading("6.1 HIS / External Integration", level=2)
    add_bullets(doc, [
        "HospitalsController: GET pulls orders from ExternalAPIBaseUri + TestRequestUri",
        "ExternalApiManager.SaveHISTestDetails: creates patient, test request, parameters",
        "ExternalApiManager.SubmitHISTestResult: outbound JSON to TestResultUri",
        "DoctorReview inline HIS submit commented — noted as SQL job in code",
        "GlobalScheduler pings external API on interval (SchedulerIntervalMinute)",
    ])

    # 7 Reports section detail
    doc.add_heading("7. Reporting Capability", level=1)
    add_bullets(doc, [
        "Operational registers with ApiOption header filters (date, patient, doctor, invoice no)",
        "Maximum report date span: 2 calendar months (ReportPageBase)",
        "Excel export via report-excel-export.util.ts",
        "Diagnostic report: gender/age-aware reference ranges from HISParameterRangMaster",
        "Abnormal flags H/L from numeric min/max comparison",
        "A4 portrait print CSS with letterhead margin, hides app chrome",
        "Lab No dropdown lists only printable cases (paid, approved, results present)",
    ])

    # 8 Security
    doc.add_heading("8. Security & Access Control", level=1)
    add_table(doc, ["Mechanism", "Implementation"], [
        ("Authentication", "OAuth 2.0 bearer tokens via OWIN; POST /Token; 14-day expiry"),
        ("Client identification", "accesskey HTTP header on token request"),
        ("Authorization", "QAuthorize: module name + bitwise permission flags"),
        ("Administrator", "Role bypasses all module checks"),
        ("Permissions", "CanAdd=1, CanEdit=2, CanAuthorize=4, CanReject=8, CanDelete=16, CanView=32"),
        ("UI enforcement", "Left-nav hasAccess/hasReportAccess mirrors API modules"),
        ("Analyzer API", "LisController AllowAnonymous (HMAC noted as TODO in code)"),
    ], [1.5, 3.0])

    # 9 Strengths
    doc.add_heading("9. UI/UX & Technical Strengths", level=1)
    add_bullets(doc, [
        "Zorya branded theme with CSS variables and consistent report styling",
        "Reusable list-module reduces CRUD screen duplication across 20+ entities",
        "Single master-form pattern for most setup/master entities",
        "Configurable MASTER_LIST_SCHEMAS drives columns and API without per-screen boilerplate",
        "Separated API/portal deployment supports independent scaling and updates",
        "Manager/repository layering keeps business rules out of controllers",
        "Effective-dated rate engine supports real laboratory pricing contracts",
        "End-to-end workflow from invoice to authorized diagnostic print",
    ])

    # 10 Deployment
    doc.add_heading("10. Deployment & Scalability", level=1)
    add_bullets(doc, [
        "On-premise IIS deployment documented under I:\\Projects\\PROD\\AVILIS",
        "MSBuild Release for API; Angular production build with legacy OpenSSL for Node 22",
        "Robocopy deploy preserves Web.config and portal web.config template",
        "SQL grant script for IIS app pool database access",
        "Multi-location readiness: separate API URL patched into portal build per environment",
        "Cloud suitability: standard .NET/IIS/SQL stack; no cloud-specific code required",
        "Desktop utilities (LisTCPIPConsole, BarcodePrint) deploy separately per workstation",
    ])

    # 11 Commercial
    doc.add_heading("11. Commercial Positioning Inputs (Evidence-Based)", level=1)
    add_table(doc, ["Positioning Theme", "Implemented Evidence"], [
        ("Enterprise-ready LIS architecture", ".NET 4.8 API + Angular SPA + SQL Server + modular managers"),
        ("Analyzer-ready laboratory", "HL7 TCP Metis6000 integration, equipment mapping, heartbeat monitoring"),
        ("Configurable pricing engine", "5 rate types with date-effective priority resolution"),
        ("Complete approval workflow", "Technician and doctor review with status machine"),
        ("Integrated billing", "Sale invoice with payment gate on final reports"),
        ("Operational visibility", "Invoice and booking registers with Excel export"),
        ("Professional reporting", "Diagnostic report with ranges, abnormal flags, A4 print"),
        ("Role-based security", "Module permissions per user with Administrator override"),
        ("Brandable modern UI", "Zorya theme, responsive layout patterns, ng-select lookups"),
    ], [1.8, 2.6])

    # 12 Conclusion
    doc.add_heading("12. Conclusion", level=1)
    doc.add_paragraph(
        "AVILIS is a fully implemented, deployable laboratory information system covering master data, "
        "transactions, sample workflow, analyzer integration (HL7/Metis6000), billing, authorization, "
        "and reporting. This document reflects capabilities verified in source code, API controllers, "
        "business managers, Angular routes, deployment scripts, and database entities. Items explicitly "
        "not implemented or stubbed in code include ASTM equipment routing (throws NotImplementedException), "
        "JWT (OAuth bearer used instead), and automatic HIS result submit from doctor review (deferred to SQL job per comments)."
    )

    doc.add_heading("Appendix A — Key API Surface Summary", level=2)
    api_rows = [
        ("Masters", "api/Department, api/Unit, api/ReferralDoctor, api/Corporate, api/TestGroup, ..."),
        ("HIS Catalog", "api/HisTest, api/HisParameterMaster, api/HisParameterRangeMaster"),
        ("Billing", "api/SaleInvoice, api/TestRate"),
        ("Samples/Results", "api/NewSample, api/Sample, api/Results, api/Lis"),
        ("Reports", "api/Reports/SaleInvoiceRegister, TestBookingRegister, TestReport"),
        ("Security", "POST /Token, api/Users, api/Roles, api/Permission"),
        ("Equipment", "api/Equipments, api/Heartbeat, api/EquipmentHeartbeat"),
    ]
    add_table(doc, ["Domain", "Routes"], api_rows, [1.2, 4.0])

    doc.add_heading("Appendix B — Analysis Sources", level=2)
    add_bullets(doc, [
        "web/Lis.Web — Angular application, routes, components, services",
        "web/Lis.Api — Controllers, Startup, WebApiConfig, SimpleInjectorConfig",
        "LIS.Businesslogic — Manager classes and workflow logic",
        "LIS.DtoModel — Entities, enums, report DTOs",
        "LIS.Com.Businesslogic, LisTCPIPConsole, BarcodePrint — Analyzer integration",
        "I:\\Projects\\PROD\\AVILIS\\SCRIPTS — Deployment automation",
        "Scripts/SeedSampleData.sql — Sample data scope",
    ])

    return doc


if __name__ == "__main__":
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    document = build_document()
    document.save(OUTPUT)
    print(f"Generated: {OUTPUT}")
    print(f"Size: {os.path.getsize(OUTPUT)} bytes")
