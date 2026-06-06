# -*- coding: utf-8 -*-
"""Generate test result flow analysis document from codebase facts."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"I:\Projects\LIS\testresult.docx"


def set_cell_shading(cell, fill="E8F4F6"):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "1A7F8C")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
    doc.add_paragraph()


def heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def para(doc, text):
    doc.add_paragraph(text)


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

    # Cover
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("AVILIS / ZoryaLIS")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(26, 127, 140)
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("Test Result Lifecycle — Technical & Business Analysis")
    r2.font.size = Pt(16)
    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t3.add_run("Discovery document (current implementation only — no redesign)").italic = True
    doc.add_page_break()

    # 1 Executive Summary
    heading(doc, "1. Executive Summary")
    para(doc,
         "AVILIS stores analyzer output as relational data in TestResults and TestResultDetails. "
         "Instruments do not write to the database directly. A Windows gateway application (LisTCPIPConsole) "
         "listens on TCP, parses HL7 v2.3.1 or ASTM E1394 messages, and posts JSON to the REST API (POST api/lis) "
         "using an equipment AccessKey header. The API resolves the sample and LIS test code via TestMappingMaster, "
         "links to TestRequestDetails, persists numeric parameter values, and advances workflow status to ReportGenerated (2). "
         "Technician and doctor approvals update TestResults review/authorization metadata and TestRequestDetails.ReportStatus. "
         "The Diagnostic Test Report reads approved results, enriches with HISParameterMaster and HISParameterRangMaster, "
         "computes H/L flags at print time, and requires DoctorApproved (5) plus paid invoice.")
    bullets(doc, [
        "No separate raw-result buffer table in application code — analyzer values land in TestResultDetails.",
        "TestParameters table holds ordered parameters per booking (inflow), not analyzer numeric values.",
        "QC/control samples use ControlResults / ControlResultDetails when LISTestCode is null.",
        "H/L flags are not persisted; they are calculated in TestReportManager for printing only.",
    ])
    doc.add_page_break()

    # 2 Analyzer Integration
    heading(doc, "2. Analyzer Integration Architecture")
    heading(doc, "2.1 Entry path", 2)
    add_table(doc, ["Layer", "Component", "Role"], [
        ["Desktop host", "LisTCPIPConsole (Home.cs)", "Binds TCP; selects HL7 or ASTM driver by EquipmentType"],
        ["Protocol", "LIS.Com.Businesslogic", "TCPIPHL7Command, TCPIPASTMCommand, equipment subclasses"],
        ["HTTP client", "Lis.Api.Client.CommunicationChannel", "Adds accesskey header on all API calls"],
        ["API", "LisController", "POST api/lis (results), GET api/lis/{sampleNo} (orders), PUT api/lis/{id} (ack)"],
        ["Business", "ResultManager", "Persists TestResults + TestResultDetails"],
    ])
    heading(doc, "2.2 Protocols (as implemented)", 2)
    add_table(doc, ["Protocol", "Transport", "Equipment drivers", "Framing"], [
        ["HL7 v2.3.1", "TCP", "Metis6000_TCPIPHL7Command, ErbaH560_TCPIPHL7Command", "VT (0x0B) … FS (0x1C) + CR; segments MSH, OBR, OBX"],
        ["ASTM E1394", "TCP", "EMA200_TCPIPASTMCommand", "ENQ/ACK/STX/ETX/EOT with checksum frames"],
        ["HTTP JSON", "REST", "All console → API traffic", "Result aggregate DTO"],
        ["Serial ASTM", "RS-232", "SerialCommand base only", "InitSerialCommand throws NotImplementedException"],
    ])
    heading(doc, "2.3 Equipment identification", 2)
    bullets(doc, [
        "Runtime (console): EquipmentType enum — Metis6000, ErbaH560, EMA200; PROTOCOL_NAME HL7 or ASTM.",
        "API/DB: HTTP header accesskey maps to EquipmentMaster.AccessKey via ModuleIdentity.",
        "ResultManager sets TestResult.EquipmentId from equipment resolved by AccessKey.",
        "Test/order correlation: SampleNo + LISTestCode from parsed message → GetRequestDetails.",
    ])
    heading(doc, "2.4 Session / heartbeat", 2)
    bullets(doc, [
        "LisContext: 60-second timer → POST api/heartbeat.",
        "TCPIPHL7Command / TCPIPASTMCommand: connection heartbeat timeout sends protocol ACK.",
        "SignalR LisHub.CallHeartBeat → HeartBeatProxy in console.",
        "EquipmentHeartBeat table updated by EquipmentHeartBeatManager.",
    ])

    # 3 Processing flow
    heading(doc, "3. Result Processing Workflow (Runtime)")
    steps = [
        ("1. Booking", "SaleInvoiceManager.LinkTestRequestsToLines creates TestRequestDetail per invoice line (ReportStatus=New). SampleNo often INV-date-seq-HISTestCode."),
        ("2. Sample list", "Patient registers / sample collected; status may remain New until sent to analyzer."),
        ("3. Order download", "Analyzer queries host (HL7 QRY/DSR or ASTM Q). Console calls GET api/lis/{sampleNo}. TestRequestDetailsManager.GetBySampleNo joins TestMappingMaster (equipment AccessKey) + active requests; returns LISTestCode per mapped HISTestCode."),
        ("4. Sent to equipment", "Console PUT api/lis/{testRequestId} → UpdateStatus(SentToEquipment=1)."),
        ("5. Analyzer result", "HL7: OBR sample + OBX NM segments. Metis6000 sets LISTestCode from last OBX field[3]; each OBX becomes TestResultDetails (LISParamCode, LISParamValue, LISParamUnit). ASTM: EMA200 ParseMessageAsync builds same Result DTO."),
        ("6. API post", "LisContext.SaveTestResult → POST api/lis with accesskey. LisController logs JSON, calls ResultManager.Add."),
        ("7. Match request", "SaveTestResult: GetRequestDetails(SampleNo, LISTestCode) — mapping + equipment + status SentToEquipment or ReportGenerated."),
        ("8. Persist", "If no existing TestResult for TestRequestId+LISTestCode: insert TestResult header + TestResultDetails lines; UpdateStatus(ReportGenerated=2)."),
        ("9. Orphan path", "If no matching request: GetParameterDetails tries join by sample + param; else logs and returns 0."),
        ("10. QC path", "If LISTestCode null: SaveControlResult → ControlResults + ControlResultDetails."),
        ("11. Review UI", "Technician queue filters ReportStatus=2. GET api/Sample/{id} loads ReviewTest via GetTestResultByRequestId / GetTestRunDetails."),
        ("12. Approval", "POST api/Sample (technician) or PUT api/Sample (doctor) → ReviewProcess updates ReportStatus and TestResults ReviewedBy/AuthorizedBy."),
        ("13. Print", "GET api/Reports/TestReport?labNo= → TestReportManager validates DoctorApproved + Paid + results exist."),
    ]
    for i, (title, desc) in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.add_run(f"{title}. ").bold = True
        p.add_run(desc)
    doc.add_page_break()

    # 4 Database entities
    heading(doc, "4. Database Design & Relationships")
    heading(doc, "4.1 Core tables", 2)
    add_table(doc, ["Table", "Entity", "PK", "Key FKs", "Purpose"], [
        ["TestRequestDetails", "TestRequestDetail", "Id", "PatientId", "Test booking / workflow; ReportStatus"],
        ["TestParameters", "TestParameter", "Id", "TestRequestDetailsId", "Ordered parameter list per booking (HIS codes)"],
        ["TestResults", "TestResult", "Id", "PatientId, TestRequestId, EquipmentId", "Test-level result header"],
        ["TestResultDetails", "TestResultDetails", "Id", "TestResultId", "Parameter values: LISParamCode, LISParamValue, LISParamUnit"],
        ["TestMappingMaster", "TestMappingMaster", "Id", "EquipmentId", "Maps HISTestCode ↔ LISTestCode per analyzer"],
        ["HISParameterMaster", "HISParameterMaster", "Id", "HisTestId", "LIS/HIS parameter definitions per test"],
        ["HISParameterRangMaster", "HISParameterRangMaster", "Id", "HisParameterId", "Reference ranges (gender/age/min/max)"],
        ["HisTestMaster", "HisTestMaster", "Id", "—", "Test catalog"],
        ["EquipmentMaster", "EquipmentMaster", "Id", "—", "Analyzer registry; AccessKey"],
        ["EquipmentHeartBeat", "EquipmentHeartBeat", "—", "Equipment", "Liveness"],
        ["ControlResults", "ControlResult", "Id", "EquipmentId", "QC material results"],
        ["ControlResultDetails", "ControlResultDetails", "Id", "ControlResultId", "QC parameter values"],
        ["SaleInvoices", "SaleInvoice", "Id", "PatientId", "Billing header"],
        ["SaleInvoiceDetails", "SaleInvoiceDetail", "Id", "SaleInvoiceId, RequestDetailId", "Line links to TestRequestDetail"],
        ["PatientDetails", "PatientDetail", "Id", "—", "Patient demographics"],
    ])
    heading(doc, "4.2 Non-persisted / legacy", 2)
    bullets(doc, [
        "TestValues, TestRun, ReviewTest — in-memory DTOs for UI/report (ReviewTest.cs).",
        "TestResultsHist / TestResultDetailsHist — entity classes exist but not in ApplicationDBContext or active C# writes.",
        "Result wrapper (LIS.DtoModel.Models.Result) — API transport only.",
    ])
    heading(doc, "4.3 Unique constraint", 2)
    para(doc, "TestRequestDetails: unique index on (SampleNo, HISTestCode, ReportStatus) — allows multiple workflow rows per sample/test when status differs (e.g. re-run).")

    # 5 Result storage
    heading(doc, "5. Result Storage Design")
    heading(doc, "5.1 Where values live", 2)
    add_table(doc, ["Data type", "Storage", "Written by", "Read by"], [
        ["Analyzer numeric values", "TestResultDetails.LISParamValue", "ResultManager.Add", "TestRequestDetailsManager, TestReportManager"],
        ["Test-level metadata", "TestResults (dates, specimen, equipment)", "ResultManager.Add", "Review + report managers"],
        ["Workflow state", "TestRequestDetails.ReportStatus", "ResultManager, TestRequestDetailsManager", "Queues, TestReportManager"],
        ["Tech review audit", "TestResults.ReviewedBy, ReviewDate, TechnicianNote", "UpdateTestResulDetails", "UI, report header"],
        ["Doctor authorization", "TestResults.AuthorizedBy, AuthorizationDate", "UpdateTestResulDetails", "Diagnostic report date"],
        ["Ordered params (booking)", "TestParameters", "PatientDetailManager, ExternalApiManager, re-open", "RawParameters API, lists"],
        ["QC values", "ControlResultDetails", "ResultManager.SaveControlResult", "QualityControlManager"],
    ])
    heading(doc, "5.2 Duplicate handling", 2)
    para(doc, "SaveTestResult: if TestResult already exists for same TestRequestId + LISTestCode, returns resultId=0 and does not insert duplicate lines. Additional OBX posts for unmatched bookings may append details only when GetParameterDetails finds existing result.")

    # 6 Sample & test relationship
    heading(doc, "6. Sample & Test Relationship Flow")
    para(doc, "End-to-end chain as implemented:")
    chain = [
        "PatientMaster / registration → PatientDetails",
        "Sale Invoice (SaleInvoiceManager.Save) → SaleInvoice + SaleInvoiceDetail lines",
        "LinkTestRequestsToLines → TestRequestDetail (HISRequestNo ≈ InvoiceNo, SampleNo, HISTestCode, ReportStatus=New)",
        "SaleInvoiceDetail.RequestDetailId FK → TestRequestDetail.Id",
        "Optional: TestParameters copied per HIS parameter for the test (booking expectation)",
        "Barcode / sample collection (UI) — SampleNo on request",
        "Analyzer order: GetBySampleNo filters requests where TestMappingMaster.LISTestCode is configured for equipment",
        "Result save: match SampleNo + LISTestCode → TestRequestDetail → TestResults.TestRequestId",
        "Parameters attach via TestResultId on TestResultDetails",
    ]
    bullets(doc, chain)

    # 7 Parameter mapping
    heading(doc, "7. Parameter Mapping & Reference Ranges")
    heading(doc, "7.1 Master data roles", 2)
    add_table(doc, ["Master", "Role in analyzer flow"], [
        ["TestMappingMaster", "Per EquipmentId: HISTestCode + LISTestCode (+ descriptions). Required for GetBySampleNo and GetRequestDetails."],
        ["HISParameterMaster", "Defines HISParamCode, LISParamCode, unit, method per HisTestId. Used to display HIS names on review."],
        ["HISParameterRangMaster", "Ranges per HisParameterId: gender, age band, MinValue, MaxValue, HISRangeValue."],
        ["TestParameter (table)", "Per TestRequestDetailsId: which parameters were ordered — not analyzer values."],
    ])
    heading(doc, "7.2 Review enrichment", 2)
    para(doc, "GetTestRunDetails joins TestResultDetails with HISParameterMaster (match LISParamCode + HISTestCode via TestMappingMaster for equipment). Loads HISParameterRangMaster rows where HISRangeCode equals HISParamCode for display strings. Optional JSON file mappings per equipment model (IFileHandler.GetJsonMappings).")
    heading(doc, "7.3 Abnormal flags", 2)
    bullets(doc, [
        "Review screen: range text only in TestValues.HISRangeValues[] — no H/L column.",
        "Diagnostic print: TestReportManager.ApplyReferenceRange — numeric parse; if MinValue>0 and value<Min → L; if MaxValue>0 and value>Max → H; sets IsAbnormal on DTO only.",
        "Flags require MinValue or MaxValue > 0; zero bounds ignored for flagging.",
    ])
    para(doc, "EquipmentParamMappingManager: stub (NotImplementedException) — parameter mapping API not implemented; mapping relies on TestMappingMaster + HISParameterMaster LISParamCode alignment.")

    # 8 Approval
    heading(doc, "8. Approval Workflow")
    add_table(doc, ["Value", "Enum", "Typical trigger"], [
        ["0", "New", "Invoice booking / re-run creates new request"],
        ["1", "SentToEquipment", "PUT api/lis/{id} after order sent"],
        ["2", "ReportGenerated", "First POST api/lis result saved"],
        ["3", "TechnicianApproved", "POST api/Sample status=3"],
        ["4", "TechnicianRejected", "POST api/Sample status=4"],
        ["5", "DoctorApproved", "PUT api/Sample status=5 — required for diagnostic print"],
        ["6", "DoctorRejected", "PUT api/Sample status=6"],
        ["7", "FinallyRejected", "Re-run closes old request"],
    ])
    bullets(doc, [
        "APIs: SampleController (GET/POST/PUT api/Sample), PatientsController (GET/PUT batch), DailyStatus.",
        "Permissions: Technician — Reports module CanAuthorize/CanReject; Doctor — DoctorsApprovals module.",
        "Multi-run: ReviewProcess uses runIndex (recentTestRequestId); other runs for same SampleNo+HISTestCode may be set TechnicianRejected.",
        "HIS outbound submit after doctor approval is commented out; documented as SQL job alternative.",
    ])

    # 9 Report generation
    heading(doc, "9. Report Generation Flow")
    bullets(doc, [
        "Endpoint: GET api/Reports/TestReport?labNo= (OperationalReportsController → TestReportManager.GetDiagnosticTestReport).",
        "Resolve invoice: by InvoiceNo or LabNo (HISRequestNo on TestRequestDetail).",
        "Gates: Invoice not cancelled; PaymentStatus=Paid (2); all linked requests ReportStatus=DoctorApproved (5); TestResults + TestResultDetails exist.",
        "Data: ResolveTestRequests from SaleInvoiceDetail.RequestDetailId or HISRequestNo=InvoiceNo.",
        "Per test: GetTestResultByRequestId → approved TestRun (DoctorApproved) → MapParameter → ApplyReferenceRange.",
        "UI: test-report.component (Angular) — A4 print layout.",
        "GET api/Reports/TestReportLabNumbers — printable lab numbers (paid + has results).",
    ])

    # 10 E2E diagram (text)
    heading(doc, "10. End-to-End Data Flow (Logical)")
    flow = """Patient Registration → PatientDetails
Sale Invoice Save → SaleInvoice / SaleInvoiceDetails
  → TestRequestDetail (New) + optional TestParameters
Sample Collection → SampleNo on request
LisTCPIPConsole TCP listen → HL7/ASTM parse
  → GET api/lis/{sampleNo} (orders via TestMappingMaster)
  → PUT api/lis/{id} (SentToEquipment)
  → POST api/lis (Result DTO)
ResultManager → TestResults + TestResultDetails (ReportGenerated)
Technician POST api/Sample → TechnicianApproved (3)
Doctor PUT api/Sample → DoctorApproved (5)
GET api/Reports/TestReport → DiagnosticTestReportDto → Print"""
    para(doc, flow)

    # 11 Relationships
    heading(doc, "11. Entity Relationship Summary")
    add_table(doc, ["From", "To", "Cardinality", "Join key"], [
        ["SaleInvoice", "SaleInvoiceDetail", "1:N", "SaleInvoiceId"],
        ["SaleInvoiceDetail", "TestRequestDetail", "N:1", "RequestDetailId"],
        ["TestRequestDetail", "PatientDetail", "N:1", "PatientId"],
        ["TestRequestDetail", "TestParameter", "1:N", "TestRequestDetailsId"],
        ["TestRequestDetail", "TestResult", "1:N", "TestRequestId"],
        ["TestResult", "TestResultDetails", "1:N", "TestResultId"],
        ["TestResult", "EquipmentMaster", "N:1", "EquipmentId"],
        ["TestMappingMaster", "EquipmentMaster", "N:1", "EquipmentId"],
        ["HISParameterMaster", "HisTestMaster", "N:1", "HisTestId"],
        ["HISParameterRangMaster", "HISParameterMaster", "N:1", "HisParameterId"],
    ])

    # 12 Runtime sequence
    heading(doc, "12. Runtime Sequence (Result POST)")
    seq = [
        "Analyzer → TCP → Metis6000_TCPIPHL7Command.ResultProcess / EMA200 ParseMessageAsync",
        "Build Result { TestResult, ResultDetails[] }",
        "LisContext.SaveTestResult → HTTP POST api/lis + header accesskey",
        "ModuleIdentity reads AccessKey",
        "ResultManager.Add → SaveTestResult or SaveControlResult",
        "equipmentRepo.Get(AccessKey) → EquipmentId",
        "GetRequestDetails(SampleNo, LISTestCode)",
        "testResultRepo.Add + resultDetailsRepo.Add (loop)",
        "testRequestDetailManager.UpdateStatus(ReportGenerated)",
        "Commit via GenericUnitOfWork / ModuleRepo",
    ]
    bullets(doc, seq)

    # 13 Error handling
    heading(doc, "13. Error Handling & Data Safety")
    add_table(doc, ["Scenario", "Current behavior"], [
        ["No matching TestRequestDetail", "Logs debug; returns 0; may try GetParameterDetails append path"],
        ["Duplicate result same request", "existResult != null → skip insert (resultId=0)"],
        ["Missing TestMapping", "GetBySampleNo returns empty — analyzer gets no orders"],
        ["Missing AccessKey", "Ping fails; equipment not resolved on save"],
        ["API exception on POST", "LisController logs exception; returns 500 with message"],
        ["Partial OBX", "All OBX NM segments collected; LISTestCode from last OBX in Metis driver"],
        ["Invalid parameter mapping", "Review shows LIS codes; HIS name empty if no HISParameterMaster match"],
        ["Unauthenticated API", "Lis endpoints [AllowAnonymous] — security relies on AccessKey knowledge"],
        ["EquipmentParamMapping API", "Not implemented — stub manager"],
    ])

    # 14 APIs summary
    heading(doc, "14. APIs & Services Index")
    add_table(doc, ["API", "Manager / Service", "Function"], [
        ["POST api/lis", "ResultManager", "Save analyzer results"],
        ["GET api/lis/{sampleNo}", "TestRequestDetailsManager", "Download orders to analyzer"],
        ["PUT api/lis/{id}", "TestRequestDetailsManager", "SentToEquipment"],
        ["POST api/heartbeat", "EquipmentHeartBeatManager", "Equipment alive"],
        ["POST api/Results", "ResultManager", "Alternate result CRUD"],
        ["GET api/RawSample", "TestRequestDetailsManager", "Read bookings"],
        ["GET api/RawParameters", "TestRequestDetailsManager", "Read TestParameters"],
        ["POST/PUT api/Sample", "TestRequestDetailsManager", "Tech/doctor approval"],
        ["GET api/Reports/TestReport", "TestReportManager", "Diagnostic report DTO"],
        ["GET api/Quality", "QualityControlManager", "QC results (separate workflow)"],
    ])

    # 15 Technical observations / risks
    heading(doc, "15. Technical Observations & Gaps")
    bullets(doc, [
        "HL7 Metis6000 driver sets LISTestCode from OBX-3 (parameter code), not a separate test code field — mapping design must align.",
        "No application-level raw message archive table — only API request logging.",
        "Duplicate result posts are silently ignored when header already exists.",
        "History tables (TestResultsHist) not wired in EF.",
        "EquipmentParamMappingManager not implemented.",
        "Serial analyzer path not factory-wired (InitSerialCommand not implemented).",
        "LisController and ResultsController marked AllowAnonymous — AccessKey is sole equipment auth.",
        "Doctor note field exists on TestResult but approval appends to TechnicianNote per UpdateTestResulDetails.",
        "HIS result submit after approval disabled in code (SQL job assumed).",
        "H/L flags not stored — reprint recalculates from current range master.",
    ])

    # Final summary
    doc.add_page_break()
    heading(doc, "16. Final Summary Checklist")
    add_table(doc, ["Topic", "Mechanism"], [
        ["Analyzer entry", "LisTCPIPConsole → HL7/ASTM TCP → POST api/lis"],
        ["Equipment ID", "EquipmentMaster.AccessKey HTTP header"],
        ["Test mapping", "TestMappingMaster: HISTestCode ↔ LISTestCode per EquipmentId"],
        ["Result storage", "TestResults + TestResultDetails"],
        ["Workflow flag", "TestRequestDetails.ReportStatus (0–7)"],
        ["Approval audit", "TestResults.ReviewedBy / AuthorizedBy + dates"],
        ["Report output", "TestReportManager → DiagnosticTestReportDto; gates: Paid + DoctorApproved"],
        ["Reference ranges", "HISParameterRangMaster; H/L at print time only"],
        ["QC path", "ControlResults when LISTestCode is null"],
    ])

    para(doc, "Document generated from source analysis of repository: I:\\Projects\\LIS\\avs-lis")
    para(doc, "Assumptions (not verified in code): Any SQL triggers populating TestResultsHist; external HIS integration job for approved results.")

    doc.save(OUTPUT)
    print(f"Written: {OUTPUT}")


if __name__ == "__main__":
    build()
